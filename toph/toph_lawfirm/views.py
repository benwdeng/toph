from django.shortcuts import render, redirect, get_object_or_404
from django.http import JsonResponse
from django.core.exceptions import ObjectDoesNotExist, PermissionDenied
from django.contrib import messages
from django.contrib.auth import authenticate, login
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import AuthenticationForm
from django.contrib.auth.views import LogoutView
from django.urls import reverse_lazy
from django.contrib.messages import get_messages
from django.core import serializers
from django.views.decorators.csrf import csrf_exempt
import os
from django.db.models import F, Q
from .models import Entry, Approval, Review, Signature, Lodge, Document, Party, LawFirmCouncil, LodgeFile
from .forms import EntryForm, DocumentForm, PartyForm, LodgeForm, EditEntryForm, ReviewForm, EditReviewForm
from django.views.generic import TemplateView
from django.conf import settings

from functools import wraps

from django.http import HttpResponse
from google.cloud import storage
from docx import Document as DocxDocument
from io import BytesIO
from django.core.files.base import ContentFile

import logging
from django.views.decorators.http import require_http_methods

from docusign_esign import ApiClient, EnvelopesApi, EnvelopeDefinition, Signer, SignHere, Recipients, Tabs, Reminders, Expirations, Notification, Witness, Text
from docusign_esign import Document as DocusignDocument
from docusign_esign.client.api_exception import ApiException
import json
import base64
import requests
from pathlib import Path
import jwt
from django.core.mail import EmailMessage
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)




@login_required(login_url='toph_lawfirm_login')
def entry_list(request):
    addressfinder_api_key = os.getenv('ADDRESSFINDER_API_KEY')
    user_groups = request.user.groups.filter(name__icontains='Law firm').values_list('name', flat=True)
    entries = Entry.objects.filter(law_firm__in=user_groups)
    documents = Document.objects.all()
    council_info = LawFirmCouncil.objects.first()

    entries_active = entries.filter(status__in=('0_costing', '1_engagement', '2_drafting', '3_review', '4_signing','5_lodgement'))
    entries_closed = entries.filter(status__in=('6_closed', ''))
    
    form = EntryForm()
    form_edit = EditEntryForm()
    form_document = DocumentForm()
    form_party = PartyForm()

    if request.method == 'POST':
        if 'form_action' in request.POST and request.POST['form_action'] == 'add_entry':
            form = EntryForm(request.POST, request.FILES)
            logger.info(request.POST)
            if form.is_valid():
                instance = form.save(commit=False)
                instance.law_firm = user_groups[0]  # Assuming user_groups[0] gives the correct law firm name
                instance.save()
                # If you need to redirect to the new entry's detail page, use instance.id
                return redirect('toph_lawfirm_entry_list')
            
    return render(
        request, 
        'toph_lawfirm/entry_list.html', 
        {
            'entries': entries,
            'council_info': council_info,
            'documents': documents,
            'entries_active': entries_active,
            'entries_closed': entries_closed,
            'user_groups': user_groups, 
            'form': form,
            'form_edit': form_edit,
            'form_document': form_document,
            'form_party': form_party,
            'addressfinder_api_key': addressfinder_api_key,
        }
    )

@login_required(login_url='toph_lawfirm_login')
def edit_entry(request, pk):
    entry = get_object_or_404(Entry, pk=pk)
    if request.method == "POST":
        form = EditEntryForm(request.POST, request.FILES, instance=entry)
        if form.is_valid():
            form.save()
            return JsonResponse({'success': True})
        else:
            return JsonResponse({'success': False, 'errors': form.errors}, status=400)
    else:
        form = EditEntryForm(instance=entry)
        return render(request, 'toph_lawfirm/edit_entry.html', {'form': form, 'entry': entry})
    
def edit_lodge(request, pk):
    lodge = get_object_or_404(Lodge, rdm_num=pk)
    if request.method == "POST":
        form = LodgeForm(request.POST, request.FILES, instance=lodge)
        if form.is_valid():
            form.save()
            return redirect('thank_you') 
        else:
            return JsonResponse({'success': False, 'errors': form.errors}, status=400)
    else:
        form = LodgeForm(instance=lodge)
        return render(request, 'toph_lawfirm/edit_lodge.html', {'form': form, 'lodge': lodge})
   
def edit_review(request, pk):
    review = get_object_or_404(Review, rdm_num=pk)
    approval = review.approval
    entry = approval.entry
    if request.method == "POST":
        form = EditReviewForm(request.POST, request.FILES, instance=review)
        if form.is_valid():
            form.save()
            return redirect('thank_you') 
        else:
            return JsonResponse({'success': False, 'errors': form.errors}, status=400)
    else:
        form = EditReviewForm(instance=review)
        return render(request, 'toph_lawfirm/edit_review.html', {'form': form, 'review': review, 'entry': entry})

@csrf_exempt
@login_required(login_url='toph_lawfirm_login')
def set_approval_outcome(request):
    if request.method == "POST":
        pk = request.POST.get('approval_id')
        approval = get_object_or_404(Approval, pk=pk)
        approval.outcome = request.POST.get('outcome')
        approval.save()
        return JsonResponse({'success': True})

    return JsonResponse({'error': 'Invalid request'}, status=400)

@login_required(login_url='toph_lawfirm_login')
def entry_detail(request, pk):
    entry = get_object_or_404(Entry, pk=pk)
    return render(request, 'toph_lawfirm/entry_detail.html', {'entry': entry})

@login_required(login_url='toph_lawfirm_login')
def edit_council_info(request):
    logger.info(request.POST)
    info = LawFirmCouncil.objects.get(pk=request.POST.get('lawfirmcouncilid'))
    info.council_email = request.POST.get('email')
    info.council_phone = request.POST.get('phone')
    info.council_address = request.POST.get('address')
    info.save()
    return JsonResponse({'success': True})  # Redirect to your entries list view

@login_required(login_url='toph_lawfirm_login')
def delete_entry(request, pk):
    entry = Entry.objects.get(pk=pk)
    entry.delete()
    return redirect('toph_lawfirm_entry_list')  # Redirect to your entries list view

@login_required(login_url='toph_lawfirm_login')
def delete_document(request):
    document = Document.objects.get(pk=pk)
    document.delete()
    return JsonResponse({'success': True})

@login_required(login_url='toph_lawfirm_login')
def delete_party(request, pk):
    party = Party.objects.get(pk=pk)
    party.delete()
    return JsonResponse({'success': True})

@login_required(login_url='toph_lawfirm_login')
def delete_approval(request, pk):
    approval = Approval.objects.get(pk=pk)
    approval.delete()
    return JsonResponse({'success': True})

@login_required(login_url='toph_lawfirm_login')
def entry_details(request, pk):
    entry = get_object_or_404(Entry, pk=pk)
    approvals = Approval.objects.filter(entry=entry)
    signatures = Signature.objects.filter(entry=entry)

    approvals_list = []
    signatures_list = []
    lodge_list = []
    party_list = []
    for approval in approvals:
        # Fetch reviews for the current approval
        reviews = approval.review_set.all()  # Assuming 'review_set' is the related name from Approval to Review
        reviews_list = [{
            'party': review.party.pk if review.party else None,
            'party_name': review.party.name if review.party else "Unknown",
            'date_sent': review.date_sent,
            'date_reviewed': review.date_reviewed,
            'comment': review.comment,
            'document': review.document.url if review.document else None,
        } for review in reviews]

        # Append approval and its reviews to the approvals list
        approvals_list.append({
            'pk': approval.pk,
            'date': approval.date,
            'document': approval.document.url if approval.document else None,
            'outcome': approval.outcome,
            'reviews_count': approval.reviews_count,
            'reviewed_count': approval.reviewed_count,
            'reviews': reviews_list,  # Include reviews in the approval object
        })

    for signature in signatures:
        # Fetch reviews for the current approval
        signers = signature.signer_set.all()  # Assuming 'review_set' is the related name from Approval to Review
        signers_list = [{
            'person': signer.person,
            'date': signer.date,
            'status': signer.status
        } for signer in signers]

        # Append approval and its reviews to the approvals list
        signatures_list.append({
            'pk': approval.pk,
            'document': signature.document.url if signature.document else None,
            'date': signature.date,  
            'signers': signers_list
        })

    lodge_list = [{
        'pk': lodge.pk,
        'to': lodge.to,
        'date': lodge.date,
        'dealing_number': lodge.dealing_number,
        'updated_cert': lodge.updated_cert.url if lodge.updated_cert else None,
    } for lodge in Lodge.objects.filter(entry=entry)]

    party_list = [{
        'pk': party.pk,
        'name': party.name,
        'role': party.role,
        'email': party.email,
        'address': party.address
    } for party in Party.objects.filter(entry=entry)]

    return JsonResponse({
        'pk': entry.pk,
        'rdm_num': entry.rdm_num,
        'law_firm': entry.law_firm,
        'council': entry.council,
        'subject_land_address': entry.subject_land_address,
        'created_date': entry.created_date,
        'last_updated': entry.last_updated,
        'status': entry.get_status_display(),
        'approvals': approvals_list,  # Include approvals with their reviews
        'signature': signatures_list,  
        'lodge': lodge_list,  
        'party': party_list,
    })

def custom_login(request):
    # Clear previous messages
    storage = get_messages(request)
    for message in storage:
        # Optionally do something with the message
        pass
    storage.used = True  # Clear messages

    if request.method == 'POST':
        form = AuthenticationForm(request, data=request.POST)
        if form.is_valid():
            username = form.cleaned_data.get('username')
            password = form.cleaned_data.get('password')
            user = authenticate(username=username, password=password)
            if user is not None:
                login(request, user)
                return redirect('toph_lawfirm_select_council')  # Adjust the redirect as needed
            else:
                messages.error(request, 'Invalid username or password.')
    else:
        form = AuthenticationForm()
    return render(request, 'toph_lawfirm/login.html', {'form': form})


class custom_logout(LogoutView):
    next_page = reverse_lazy('landing_page')  

@login_required(login_url='toph_lawfirm_login')
def select_council(request):
    return render(request, 'toph_lawfirm/select_council.html')   



@login_required(login_url='toph_lawfirm_login')
def handle_draft_form(request):
    if request.method == "POST":
        # Extract form data
        fields = [
            "entryId",
            "SubjectLand",
            "draftOwner1",
            "draftOwner2",
            "draftOwner3",
            "draftOwner4",
            "draftOwner5",
            "draftOwner6",
            "draftOwner7",
            "CoTVolumeNumber", "CoTFolioNumber",
            "DevelopmentPermitNumber", "DevelopmentPermitIssueDate", "DevelopmentPermitCondition", 
            "SubdivisionPermitNumber", "SubdivisionPermitIssueDate", "SubdivisionPermitCondition",
            "CouncilSignoffClause",
            "EntryMethod",
        ]


        # Use dictionary comprehension to extract these fields from request.POST
        replacements = {field: request.POST.get(field, "") for field in fields}

        doc_content = None
        file_name = "Draft Section 173 Agreement - " + replacements["SubjectLand"] + '.docx'

        if replacements['EntryMethod'] == 'upload':
            uploaded_file = request.FILES.get('UploadDocument')
            if uploaded_file:
                # Read the content of the uploaded file
                doc_content = uploaded_file.read()
                file_name = uploaded_file.name 
        else:


            replacements['Owner1Name'] = Party.objects.get(pk=replacements['draftOwner1']).name if replacements['draftOwner1'] != "" else replacements.get('Owner1Name', '')
            replacements['Owner2Name'] = Party.objects.get(pk=replacements['draftOwner2']).name if replacements['draftOwner2'] != "" else replacements.get('Owner2Name', '')
            replacements['Owner3Name'] = Party.objects.get(pk=replacements['draftOwner3']).name if replacements['draftOwner3'] != "" else replacements.get('Owner3Name', '')
            replacements['Owner4Name'] = Party.objects.get(pk=replacements['draftOwner4']).name if replacements['draftOwner4'] != "" else replacements.get('Owner4Name', '')
            replacements['Owner5Name'] = Party.objects.get(pk=replacements['draftOwner5']).name if replacements['draftOwner5'] != "" else replacements.get('Owner5Name', '')
            replacements['Owner6Name'] = Party.objects.get(pk=replacements['draftOwner6']).name if replacements['draftOwner6'] != "" else replacements.get('Owner6Name', '')
            replacements['Owner7Name'] = Party.objects.get(pk=replacements['draftOwner7']).name if replacements['draftOwner7'] != "" else replacements.get('Owner7Name', '')

            replacements['Owner1Address'] = Party.objects.get(pk=replacements['draftOwner1']).address if replacements['draftOwner1'] != "" else replacements.get('Owner1Address', '')
            replacements['Owner2Address'] = Party.objects.get(pk=replacements['draftOwner2']).address if replacements['draftOwner2'] != "" else replacements.get('Owner2Address', '')
            replacements['Owner3Address'] = Party.objects.get(pk=replacements['draftOwner3']).address if replacements['draftOwner3'] != "" else replacements.get('Owner3Address', '')
            replacements['Owner4Address'] = Party.objects.get(pk=replacements['draftOwner4']).address if replacements['draftOwner4'] != "" else replacements.get('Owner4Address', '')
            replacements['Owner5Address'] = Party.objects.get(pk=replacements['draftOwner5']).address if replacements['draftOwner5'] != "" else replacements.get('Owner5Address', '')

            replacements['Owner1Email'] = Party.objects.get(pk=replacements['draftOwner1']).email if replacements['draftOwner1'] != "" else replacements.get('Owner1Address', '')

            replacements["CouncilName"] = replacements["Owner1Name"]
            replacements["CouncilAddress"] = replacements["Owner1Address"]
            replacements["CouncilEmail"] = replacements["Owner1Email"]

            replacements["Owner1FinalName"] = replacements["Owner2Name"]
            replacements["Owner2FinalName"] = replacements["Owner3Name"]
            replacements["Owner3FinalName"] = replacements["Owner4Name"]
            replacements["Owner4FinalName"] = replacements["Owner5Name"]
            replacements["Owner5FinalName"] = ''

            replacements["Owner1Name"] = replacements["Owner1FinalName"]
            replacements["Owner2Name"] = replacements["Owner2FinalName"]
            replacements["Owner3Name"] = replacements["Owner3FinalName"]
            replacements["Owner4Name"] = replacements["Owner4FinalName"]
            replacements["Owner5Name"] = ''

            replacements["Owner1Address"] = replacements['Owner2Address']
            replacements["Owner2Address"] = replacements['Owner3Address']
            replacements["Owner3Address"] = replacements['Owner4Address']
            replacements["Owner4Address"] = replacements['Owner5Address']
            replacements["Owner5Address"] = ''

            replacements["Owner1ACN"] = ""
            replacements["Owner2ACN"] = ""
            replacements["Owner3ACN"] = ""
            replacements["Owner4ACN"] = ""
            replacements["Owner5ACN"] = ""
            replacements["Owner1NumberDirectors"] = "1"
            replacements["Owner2NumberDirectors"] = "1"
            replacements["Owner3NumberDirectors"] = "1"
            replacements["Owner4NumberDirectors"] = "1"
            replacements["Owner5NumberDirectors"] = "1"

            replacements["MortgageInstitution"] = replacements['Owner6Name']
            replacements["CaveatorName"] = replacements['Owner7Name']


            # Initialize GCS client and define bucket and file name
            client = storage.Client()
            bucket_name = "toph_lawfirm"
            template_file_name = "inputs/Section 173 Agreement - template (14).docx"
            bucket = client.bucket(bucket_name)
            blob = bucket.blob(template_file_name)

            # Download the template from GCS
            template_content = blob.download_as_bytes()
            doc = DocxDocument(BytesIO(template_content))


            # Identify the table you want to delete
            tables_to_delete = []
            if replacements["Owner1FinalName"] == "":
                tables_to_delete += [0,6]
            if replacements["Owner1Name"] == "":
                tables_to_delete += [0,6,12]
            if replacements["Owner1ACN"] == "":
                tables_to_delete += [13,14]
            if replacements["Owner1NumberDirectors"] == "1":
                tables_to_delete += [13]
            if replacements["Owner1NumberDirectors"] == "2":
                tables_to_delete += [14]

            if replacements["Owner2FinalName"] == "":
                tables_to_delete += [1,7]
            if replacements["Owner2Name"] == "":
                tables_to_delete += [15]
            if replacements["Owner2ACN"] == "":
                tables_to_delete += [16,17]
            if replacements["Owner2NumberDirectors"] == "1":
                tables_to_delete += [16]
            if replacements["Owner2NumberDirectors"] == "2":
                tables_to_delete += [17]

            if replacements["Owner3FinalName"] == "":
                tables_to_delete += [2,8]
            if replacements["Owner3Name"] == "":
                tables_to_delete += [18]
            if replacements["Owner3ACN"] == "":
                tables_to_delete += [19,20]
            if replacements["Owner3NumberDirectors"] == "1":
                tables_to_delete += [19]
            if replacements["Owner3NumberDirectors"] == "2":
                tables_to_delete += [20]

            if replacements["Owner4FinalName"] == "":
                tables_to_delete += [3,9]
            if replacements["Owner4Name"] == "":
                tables_to_delete += [21]
            if replacements["Owner4ACN"] == "":
                tables_to_delete += [22,23]
            if replacements["Owner4NumberDirectors"] == "1":
                tables_to_delete += [22]
            if replacements["Owner4NumberDirectors"] == "2":
                tables_to_delete += [23]

            if replacements["Owner5FinalName"] == "":
                tables_to_delete += [4,10]
            if replacements["Owner5Name"] == "":
                tables_to_delete += [24]
            if replacements["Owner5ACN"] == "":
                tables_to_delete += [25,26]
            if replacements["Owner5NumberDirectors"] == "1":
                tables_to_delete += [25]
            if replacements["Owner5NumberDirectors"] == "2":
                tables_to_delete += [26]

            if replacements["MortgageInstitution"] == "":
                tables_to_delete += [27]
            if replacements["CaveatorName"] == "":
                tables_to_delete += [28]

            tables_to_delete = sorted(set(tables_to_delete), reverse=True)  # Make unique and sort in reverse order

            for i in tables_to_delete:
                table_to_delete = doc.tables[i]
                parent_element = table_to_delete._element.getparent()
                parent_element.remove(table_to_delete._element)

            # Iterate through each paragraph in the document
            for paragraph in doc.paragraphs:
                for placeholder, replacement in replacements.items():
                    if placeholder in paragraph.text:
                        for run in paragraph.runs:
                            run.text = run.text.replace(placeholder, replacement)

            # Now, handle replacements in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, replacement in replacements.items():
                                if placeholder in paragraph.text:
                                    for run in paragraph.runs:
                                        run.text = run.text.replace(placeholder, replacement)

            # Save the modified document to a BytesIO object
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_content = doc_io.getvalue()  # Read the content once

            # Reset the pointer of doc_io if needed for other operations
            doc_io.seek(0)

        entry = Entry.objects.get(id=replacements["entryId"])
        entry.status = '2_drafting'
        entry.save()

        approval = Approval(
            entry=entry,
            document=ContentFile(doc_content, name=file_name)
        )
        approval.save()

        return JsonResponse({'success': True})


@csrf_exempt
@require_http_methods(["POST"])
@login_required(login_url='toph_lawfirm_login')
def send_document_for_signing(request):
    # Load your integration key, user's access token, and other credentials



    INTEGRATION_KEY = '6b12d9ac-1ed3-43da-90ab-a7c2b514df81'
    USER_ID = 'c9297943-0b9b-487c-aaea-62621d8ddf3c'
    ACCOUNT_ID = '47725e57-bda2-4652-93aa-23b9a90c98a6'
    BASE_PATH = 'https://demo.docusign.net/restapi'
    OAUTH_BASE_URL = 'account-d.docusign.com'  # Use account.docusign.com for production

    try:
        private_key_file = settings.PRIVATE_KEY_PATH
        with private_key_file.open("r") as file:
            rsa_private_key = file.read()
    except Exception as e:
        logger.error(f"Failed to read RSA private key file: {e}")
        return JsonResponse({"error": "Internal server error"}, status=500)


    # Initialize the API client
    api_client = ApiClient()
    api_client.set_base_path(OAUTH_BASE_URL)
    api_client.set_oauth_host_name(OAUTH_BASE_URL)

    # Obtain an access token using JWT grant
    try:
        token_response = api_client.request_jwt_user_token(client_id=INTEGRATION_KEY,
                                          user_id=USER_ID,
                                          oauth_host_name=OAUTH_BASE_URL,
                                          private_key_bytes=rsa_private_key.encode('utf-8'),
                                          expires_in=3600,
                                          scopes=["signature", "impersonation"])
        access_token = token_response.access_token
        logger.info("Successfully obtained JWT token")
    except ApiException as e:
        logger.error(f"Exception when requesting token: {e}")
        return JsonResponse({"error": "Failed to authenticate with DocuSign"}, status=500)

    
    # Initialize the DocuSign API client
    api_client = ApiClient()
    api_client.host = BASE_PATH
    api_client.set_default_header("Authorization", f"Bearer {access_token}")

    logger.info(f"Incoming request data: {request.body}")
    logger.info(f"Incoming headers: {request.headers}")

    json_data = request.POST.get('json_data', '{}')
    try:
        data = json.loads(json_data)
    except json.JSONDecodeError:
        return JsonResponse({"error": "Invalid JSON format"}, status=400)

    # Now, you can access your JSON data as needed
    OwnerName1 = data.get('OwnerName1')
    OwnerName2 = data.get('OwnerName2')
    OwnerName3 = data.get('OwnerName3')
    OwnerName4 = data.get('OwnerName4')
    OwnerName5 = data.get('OwnerName5')
    OwnerOrder1 = data.get('OwnerOrder1')
    OwnerOrder2 = data.get('OwnerOrder2')
    OwnerOrder3 = data.get('OwnerOrder3')
    OwnerOrder4 = data.get('OwnerOrder4')
    OwnerOrder5 = data.get('OwnerOrder5')
    reminder_frequency = data.get('reminderFrequency')
    reminder_delay = data.get('reminderDelay')


    council_name = Party.objects.get(pk=OwnerName1).name if OwnerName1 != "" else ''
    signer1_name = Party.objects.get(pk=OwnerName2).name if OwnerName2 != "" else ''
    signer2_name = Party.objects.get(pk=OwnerName3).name if OwnerName3 != "" else ''
    signer3_name = Party.objects.get(pk=OwnerName4).name if OwnerName4 != "" else ''
    signer4_name = Party.objects.get(pk=OwnerName5).name if OwnerName5 != "" else ''

    council_email = Party.objects.get(pk=OwnerName1).email if OwnerName1 != "" else ''
    signer1_email = Party.objects.get(pk=OwnerName2).email if OwnerName2 != "" else ''
    signer2_email = Party.objects.get(pk=OwnerName3).email if OwnerName3 != "" else ''
    signer3_email = Party.objects.get(pk=OwnerName4).email if OwnerName4 != "" else ''
    signer4_email = Party.objects.get(pk=OwnerName5).email if OwnerName5 != "" else ''

    council_order = OwnerOrder1
    signer1_order = OwnerOrder2
    signer2_order = OwnerOrder3
    signer3_order = OwnerOrder4
    signer4_order = OwnerOrder5

    # For the file
    if 'documentUrl' in request.FILES:
        document_file = request.FILES['documentUrl']
        document_base64 = base64.b64encode(document_file.read()).decode('utf-8')
        # Process the file as needed
    else:
        logger.error("No document uploaded")
        return JsonResponse({"error": "No document uploaded"}, status=400)

    entry_id = request.POST.get('docusignEntryId')  # Make sure 'entryId' matches the name attribute in your form


    # Create an envelope to be signed
    envelope_definition = EnvelopeDefinition(email_subject="Please Sign this Document")
    doc = DocusignDocument(document_base64=document_base64,  # You need to convert your document to base64
                   name="Sample Document", document_id="1", file_extension="docx")
    envelope_definition.documents = [doc]

    signers_list = []
    witnesses_list = []

    if signer1_name!="":
        # Create a signer recipient to sign the document
        Signer1 = Signer(
            email=signer1_email,
            name=signer1_name,
            recipient_id="1",
            routing_order=signer1_order,
            # Specify that this signer can add a witness during the signing process
            witness={'witness_for': '1'}
        )
        Signer1SignHere = SignHere(
            document_id="1", 
            page_number="11", 
            recipient_id="1", 
            x_position="280", 
            y_position="120",
        )
        Signer1WitnessSignHere = SignHere(
            document_id="1",
            page_number="11",
            x_position="60",
            y_position="120",
        )
        Signer1WitnessName = Text(
            document_id="1",
            page_number="11",
            x_position="60",
            y_position="190",
            tab_label="Witness_name",
            required="true",
            tooltip="Enter your full name"
        )
        Signer1WitnessAddress = Text(
            document_id="1",
            page_number="11",
            x_position="60",
            y_position="240",
            tab_label="Witness_address",
            required="true",
            tooltip="Enter your address"
        )
        Signer1Witness = Witness(
            email='witness@email.com',
            name='Witness name',
            routing_order=signer1_order,
            witness_for='1', 		# Must match the signer’s recipient_id
            recipient_id='11'
        )

        Signer1Witness.tabs = Tabs(sign_here_tabs=[Signer1WitnessSignHere], text_tabs=[Signer1WitnessName, Signer1WitnessAddress])
        Signer1.tabs = Tabs(sign_here_tabs=[Signer1SignHere])
        signers_list += [Signer1]
        witnesses_list += [Signer1Witness]

    if signer2_name!="":
        # Create a signer recipient to sign the document
        Signer2 = Signer(
            email=signer2_email,
            name=signer2_name,
            recipient_id="2",
            routing_order=signer1_order,
            # Specify that this signer can add a witness during the signing process
            witness={'witness_for': '2'}
        )
        Signer2SignHere = SignHere(
            document_id="1", 
            page_number="12", 
            recipient_id="2", 
            x_position="280", 
            y_position="120",
        )
        Signer2WitnessSignHere = SignHere(
            document_id="1",
            page_number="12",
            x_position="60",
            y_position="120",
        )
        Signer2WitnessName = Text(
            document_id="1",
            page_number="12",
            x_position="60",
            y_position="190",
            tab_label="Witness_name",
            required="true",
            tooltip="Enter your full name"
        )
        Signer2WitnessAddress = Text(
            document_id="1",
            page_number="12",
            x_position="60",
            y_position="240",
            tab_label="Witness_address",
            required="true",
            tooltip="Enter your address"
        )
        Signer2Witness = Witness(
            email='witness@email.com',
            name='Witness name',
            routing_order=signer2_order,
            witness_for='2', 		# Must match the signer’s recipient_id
            recipient_id='12'
        )
        Signer2Witness.tabs = Tabs(sign_here_tabs=[Signer2WitnessSignHere], text_tabs=[Signer2WitnessName, Signer2WitnessAddress])
        Signer2.tabs = Tabs(sign_here_tabs=[Signer2SignHere])
        signers_list += [Signer2]
        witnesses_list += [Signer2Witness]

    if signer3_name!="":
        # Create a signer recipient to sign the document
        Signer3 = Signer(
            email=signer3_email,
            name=signer3_name,
            recipient_id="3",
            routing_order=signer1_order,
            # Specify that this signer can add a witness during the signing process
            witness={'witness_for': '3'}
        )
        Signer3SignHere = SignHere(
            document_id="1", 
            page_number="13", 
            recipient_id="2", 
            x_position="280", 
            y_position="120",
        )
        Signer3WitnessSignHere = SignHere(
            document_id="1",
            page_number="13",
            x_position="60",
            y_position="120",
        )
        Signer3WitnessName = Text(
            document_id="1",
            page_number="13",
            x_position="60",
            y_position="190",
            tab_label="Witness_name",
            required="true",
            tooltip="Enter your full name"
        )
        Signer3WitnessAddress = Text(
            document_id="1",
            page_number="13",
            x_position="60",
            y_position="240",
            tab_label="Witness_address",
            required="true",
            tooltip="Enter your address"
        )
        Signer3Witness = Witness(
            email='witness@email.com',
            name='Witness name',
            routing_order=signer3_order,
            witness_for='3', 		# Must match the signer’s recipient_id
            recipient_id='13'
        )
        Signer3Witness.tabs = Tabs(sign_here_tabs=[Signer3WitnessSignHere], text_tabs=[Signer3WitnessName, Signer3WitnessAddress])
        Signer3.tabs = Tabs(sign_here_tabs=[Signer3SignHere])
        signers_list += [Signer3]
        witnesses_list += [Signer3Witness]

    if signer4_name!="":
        # Create a signer recipient to sign the document
        Signer4 = Signer(
            email=signer4_email,
            name=signer4_name,
            recipient_id="4",
            routing_order=signer1_order,
            # Specify that this signer can add a witness during the signing process
            witness={'witness_for': '4'}
        )
        Signer4SignHere = SignHere(
            document_id="1", 
            page_number="14", 
            recipient_id="2", 
            x_position="280", 
            y_position="120",
        )
        Signer4WitnessSignHere = SignHere(
            document_id="1",
            page_number="14",
            x_position="60",
            y_position="120",
        )
        Signer4WitnessName = Text(
            document_id="1",
            page_number="14",
            x_position="60",
            y_position="190",
            tab_label="Witness_name",
            required="true",
            tooltip="Enter your full name"
        )
        Signer4WitnessAddress = Text(
            document_id="1",
            page_number="14",
            x_position="60",
            y_position="240",
            tab_label="Witness_address",
            required="true",
            tooltip="Enter your address"
        )
        Signer4Witness = Witness(
            email='witness@email.com',
            name='Witness name',
            routing_order=signer4_order,
            witness_for='4', 		# Must match the signer’s recipient_id
            recipient_id='14'
        )
        Signer4Witness.tabs = Tabs(sign_here_tabs=[Signer4WitnessSignHere], text_tabs=[Signer4WitnessName, Signer4WitnessAddress])
        Signer4.tabs = Tabs(sign_here_tabs=[Signer4SignHere])
        signers_list += [Signer4]
        witnesses_list += [Signer4Witness]



    SignerCouncil = Signer(
        email=council_email, 
        name=council_name, 
        recipient_id="9", 
        routing_order=council_order
    )
    SignerCouncilSignHere = SignHere(
        document_id="1", 
        page_number="10", 
        recipient_id="9", 
        x_position="280", 
        y_position="240"
    )
    SignerCouncilName = Text(
        document_id="1",
        page_number="10",
        x_position="280",
        y_position="310",
        tab_label="Name",
        required="true",
        tooltip="Enter your full name"
    )
    SignerCouncilWitness = Witness(
        email='witness@email.com',
        name='Witness name',
        routing_order=council_order,
        witness_for='9', 		# Must match the signer’s recipient_id
        recipient_id='19'
    )
    SignerCouncilWitnessSignHere = SignHere(
        document_id="1",
        page_number="10",
        x_position="60",
        y_position="240",
    )
    SignerCouncilWitnessName = Text(
        document_id="1",
        page_number="10",
        x_position="60",
        y_position="310",
        tab_label="Witness_name",
        required="true",
        tooltip="Enter your full name"
    )
    SignerCouncilWitness.tabs = Tabs(sign_here_tabs=[SignerCouncilWitnessSignHere], text_tabs=[SignerCouncilWitnessName])
    SignerCouncil.tabs = Tabs(sign_here_tabs=[SignerCouncilSignHere], text_tabs=[SignerCouncilName])
    signers_list += [SignerCouncil]
    witnesses_list += [SignerCouncilWitness]

    # Add the signer to the envelope's recipients
    envelope_definition.recipients = Recipients(signers=signers_list, witnesses=witnesses_list)
    envelope_definition.status = "sent"  # Set to "sent" to send the envelope immediately
    reminders = Reminders(reminder_enabled="true", reminder_delay=reminder_delay, reminder_frequency=reminder_frequency)
    expirations = Expirations(expire_enabled="true", expire_after="10", expire_warn="2")
    notification = Notification(reminders=reminders, expirations=expirations)
    envelope_definition.notification = notification
    try:
        # Use the EnvelopesApi to create and send the envelope
        envelopes_api = EnvelopesApi(api_client)
        envelope_summary = envelopes_api.create_envelope(account_id=ACCOUNT_ID, envelope_definition=envelope_definition)

        entry = Entry.objects.get(pk=entry_id)
        entry.status = '4_signing'
        entry.save()

        signature = Signature(
            entry=entry,
            document=ContentFile(document_file.read(), name=document_file.name)
        )
        signature.save()
        
        return JsonResponse({"status": "success", "envelope_id": envelope_summary.envelope_id}, status=200)
    except ApiException as e:
        print("API exception occurred!")
        # Print the status code if available
        if e.status:
            print(f"Status code: {e.status}")
        # Print the full error message
        print(f"Error message: {e.body}")
        # Return a more detailed error response
        return JsonResponse({"status": "error", "error": e.body.decode('utf-8') if isinstance(e.body, bytes) else e.body}, status=e.status if e.status else 400)





@csrf_exempt
@login_required(login_url='toph_lawfirm_login')
def send_lodge_email(request):
    if request.method == 'POST':
        try:

            entry_id = request.POST.get('entryId')  # Make sure 'entryId' matches the name attribute in your form

            # Find the corresponding Entry instance
            entry = Entry.objects.get(rdm_num=entry_id)

            target_email = request.POST.get('targetEmail')
            cc_email = request.POST.getlist('ccEmail')  # Can handle multiple CC emails
            email_subject = request.POST.get('emailSubject')
            email_body = request.POST.get('emailBody')
            attachments = request.FILES.getlist('emailAttachments')  # Can handle multiple files

            lodge = Lodge(
                entry=entry,
                to=target_email
            )
            lodge.save()

            for attachment in attachments:
                lodge_file = LodgeFile(lodge=lodge, file=attachment)
                lodge_file.save()

            email_body += f"\n\n Please provide the dealing number via {settings.SITE_URL}/toph_lawfirm/lodge/edit/{lodge.rdm_num}/ \n\n This email was sent by The Online Planning Hub"


            email = EmailMessage(
                subject=email_subject,
                body=email_body,
                from_email='theonlineplanninghub@outlook.com',
                to=[target_email],
                cc=cc_email if cc_email else None  # Add CC only if provided
            )

            # Attach each file
            for doc in attachments:
                email.attach(doc.name, doc.read(), doc.content_type)

            email.send()

            entry.status = '5_lodgement'
            entry.save()

            return JsonResponse({"success": True})
        except Exception as e:
            # Log the error
            print(e)
            return JsonResponse({"success": False, "error": str(e)}, status=500)

    return JsonResponse({"success": False}, status=400)




def update_dealing_number(request, pk):
    entry = get_object_or_404(Entry, pk=pk)
    if request.method == 'POST':
        form = EntryForm(request.POST, request.FILES, instance=entry)
        if form.is_valid():
            form.save()
            return redirect('thank_you')  # Redirect as appropriate
    else:
        form = EntryForm(instance=entry)

    return render(request, 'toph_lawfirm/update_dealing_number.html', {'form': form})


class ThankYouPageView(TemplateView):
    template_name = 'toph_lawfirm/thank_you.html'



@login_required(login_url='toph_lawfirm_login')
def upload_document(request, pk):
    entry = Entry.objects.get(pk=pk)  # Assuming you're passing the entry ID
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            document = form.save(commit=False)
            document.entry = entry  # Associate the uploaded document with the entry
            document.save()
            return JsonResponse({'success': True})
    else:
        form = DocumentForm()
    return JsonResponse({'success': False}, status=400)

@login_required(login_url='toph_lawfirm_login')
def upload_party(request, pk):
    entry = Entry.objects.get(pk=pk)  # Assuming you're passing the entry ID
    if request.method == 'POST':
        form = PartyForm(request.POST, request.FILES)
        if form.is_valid():
            party = form.save(commit=False)
            party.entry = entry  # Associate the uploaded document with the entry
            party.save()
            return JsonResponse({'success': True})
    else:
        form = DocumentForm()
    return JsonResponse({'success': False}, status=400)

@login_required(login_url='toph_lawfirm_login')
def add_entry(request):
    logger.info(request.POST)
    if request.method == 'POST':
        form = EntryForm(request.POST, request.FILES)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.law_firm = request.POST.get('law_firm')
            instance.council = request.POST.get('council')
            instance.save()

            party = Party(
                entry=instance,
                name=request.POST.get('council'),
                role='Council',
                email=request.POST.get('council_email'),
                phone=request.POST.get('council_phone'),
                address=request.POST.get('council_address')
            )
            party.save()
            logger.info(party)

            return JsonResponse({'success': True})

@login_required(login_url='toph_lawfirm_login')
def get_documents_for_entry(request, pk):
    documents = Document.objects.filter(entry=pk).values('pk','name', 'document')
    return JsonResponse(list(documents), safe=False)


@login_required(login_url='toph_lawfirm_login')
def get_parties_for_entry(request, pk):
    parties = list(Party.objects.filter(entry=pk).values('pk','name', 'role', 'email', 'phone', 'address'))
    return JsonResponse(parties, safe=False)

@login_required(login_url='toph_lawfirm_login')
def get_entries(request):
    user_groups = request.user.groups.filter(name__icontains='Law firm').values_list('name', flat=True)
    entries = Entry.objects.filter(law_firm__in=user_groups)
    
    # Serialize the queryset
    data = serializers.serialize('json', entries)
    data_list = json.loads(data)
    
    # Enhance each object with get_status_display()
    for entry in data_list:
        # Assuming 'fields' is where the serialized data is stored. Adjust if necessary.
        entry_obj = Entry.objects.get(pk=entry['pk'])  # Get the actual Entry object
        entry['fields']['status_display'] = entry_obj.get_status_display()  # Append the human-readable status
    
    return JsonResponse(data_list, safe=False)

@login_required(login_url='toph_lawfirm_login')
def get_reviews_for_approval(request, pk):
    reviews = Review.objects.filter(approval=pk).annotate(
        party_name=F('party__name'),
        party_pk=F('party__pk'),
        party_role=F('party__role'),
        party_email=F('party__email'),
        party_phone=F('party__phone'),
    ).values(
        'pk', 'party', 'date_sent', 'date_reviewed', 'status', 'comment', 'document', 'party_pk', 'party_name', 'party_role', 'party_email', 'party_phone'
    )
    return JsonResponse(list(reviews), safe=False)

@login_required(login_url='toph_lawfirm_login')
def upload_review(request):
    if request.method == 'POST':
        approval_id = request.POST.get('approval_id')
        party_ids = request.POST.getlist('party_id') 
        cc_email = request.POST.getlist('cc_email')
        reminder_delay = request.POST.get('reminder_delay')
        reminder_frequency = request.POST.get('reminder_frequency')
        approval = Approval.objects.get(pk=approval_id)
        entry = approval.entry
        entry.status = '3_review'
        entry.save()

        for party_id in party_ids:
            party = Party.objects.get(pk=party_id)
            new_review = Review.objects.create(approval=approval, party=party, status=None, reminder_delay=reminder_delay, reminder_frequency=reminder_frequency)

            target_email = party.email
            email_subject = 'Please Review Draft Agreement'

            if (party.role == 'Council'):
                email_body = 'Please review the draft agreement at ' + settings.SITE_URL + "/toph_council/"
            else:
                email_body = 'Please review the draft agreement at ' + settings.SITE_URL + f"/toph_lawfirm/review/edit/{new_review.rdm_num}/ \n"

            email = EmailMessage(
                subject=email_subject,
                body=email_body,
                from_email='theonlineplanninghub@outlook.com',
                to=[target_email]
            )
            email.send()

        return JsonResponse({'success': True})

    return JsonResponse({'success': False}, status=400)

from django.utils.timezone import now
from datetime import timedelta

def send_reminder_emails(request):

    # Get the current date
    current_date = now().date()

    # Fetch reviews that are potentially ready for a reminder
    # This fetches reviews that:
    # - Have not been reviewed (status is null or blank)
    # - Have a set reminder_delay and reminder_frequency
    # We apply only basic filtering here to limit the dataset.
    initial_reviews_to_remind = Review.objects.filter(
        (Q(status__isnull=True) | Q(status__exact='')),
        reminder_delay__gt=0,
        reminder_frequency__gt=0
    )
    logger.info(initial_reviews_to_remind)

    # Further filter the queryset in Python
    reviews_to_remind = []
    for review in initial_reviews_to_remind:
        if review.date_sent:
            date_sent = review.date_sent.date()
            delay_days = timedelta(days=review.reminder_delay)
            reminder_start_date = date_sent + delay_days

            # Calculate the days since the reminder period started
            days_since_reminder_start = (current_date - reminder_start_date).days

            # Check if the reminder_start_date is in the past and today is a reminder day
            if reminder_start_date <= current_date and days_since_reminder_start % review.reminder_frequency == 0:
                reviews_to_remind.append(review)

    for review in reviews_to_remind:
        if review.party:
            if review.party.email:
                logger.info(review.date_sent)
                logger.info(review.date_reviewed)
                logger.info(review.status)
                logger.info(review.reminder_delay)
                logger.info(review.reminder_frequency)
                logger.info(review.party.name)
                logger.info(review.party.email)
                # Construct your email message here
                subject = 'Review Reminder'
                message = 'Please complete your review.'
                recipient_list = [review.party.email]
                email = EmailMessage(subject, message, 'theonlineplanninghub@outlook.com', recipient_list)
                #email.send()

    return HttpResponse("Reminders triggered successfully.", status=200)

@csrf_exempt
@login_required(login_url='toph_lawfirm_login')
def share_tracking(request):
    if request.method == 'POST':
        logger.info(request.POST)
        entry_id = request.POST.get('entry_id')
        party_ids = request.POST.getlist('party_id')

        for party_id in party_ids:
            entry = Entry.objects.get(pk=entry_id)
            party = Party.objects.get(pk=party_id)

            if party.email:
                subject = f'Application for {entry.subject_land_address}'
                message = f'Application for {entry.subject_land_address} has been created. \n You can view the progress of the application via tracking number {entry.rdm_num} at {settings.SITE_URL}/toph_owner/ \n'
                recipient_list = [party.email]
                email = EmailMessage(subject, message, 'theonlineplanninghub@outlook.com', recipient_list)
                email.send()

    return HttpResponse("Shared successfully.", status=200)

@login_required(login_url='toph_lawfirm_login')
def get_documents_lodged(request, pk):
    logger.info(1)
    try:
        lodge = Lodge.objects.get(pk=pk)
        files = lodge.files.all()  # Accessing related LodgeFile instances

        files_data = [{'id': file.id, 'url': file.file.url, 'name': file.file.name.split('/')[-1]} for file in files]
        logger.info(files_data)

        return JsonResponse(files_data, safe=False)

    except Lodge.DoesNotExist:
        return JsonResponse({'error': 'Lodge not found'}, status=404)
    except ObjectDoesNotExist:
        return JsonResponse({'error': 'Invalid request'}, status=400)


